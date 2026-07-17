from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
CONTENT_PATH = Path(__file__).with_name("script_content.txt")
OUTPUT_PATH = ROOT / "outputs" / "厦大_JiaYuan_小学期课程汇报_逐字讲解稿.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "EAF3F8"
PALE_GRAY = "F2F4F7"
TEXT = "24303A"
MUTED = "5A6874"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:hint"), "eastAsia")


def set_lang(run, value="zh-CN"):
    r_pr = run._element.get_or_add_rPr()
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), value)
    lang.set(qn("w:eastAsia"), value)


def format_run(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    set_east_asia_font(run)
    set_lang(run)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "7")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    format_run(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    format_run(run2, size=9, color=MUTED)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Number", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(TEXT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.25

    if "Time Label" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Time Label", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True

    if "Stage Cue" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Stage Cue", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.name = "Calibri"
        style.font.size = Pt(10)
        style.font.italic = True
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("JiaYuan 小学期课程汇报｜逐字讲解稿")
    format_run(r, size=8.5, color=MUTED)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    add_page_number(section.footer.paragraphs[0])
    section.first_page_footer.paragraphs[0].text = ""

    doc.core_properties.title = "厦大 JiaYuan 小学期课程汇报逐字讲解稿"
    doc.core_properties.subject = "配套 14 页课程汇报 PPT 的逐字稿与答辩速查"
    doc.core_properties.author = "JiaYuan 项目组"
    doc.core_properties.keywords = "JiaYuan, RAG, FastAPI, Vue 3, Electron, 答辩讲稿"


def add_inline_emphasis(paragraph, text):
    """Bold likely code identifiers while keeping speech text readable."""
    tokens = []
    start = 0
    code_terms = (
        "backend/", "frontend/", "desktop/", ".py", ".vue", ".cjs", ".json",
        "process_and_vectorize_document_stream", "retrieve_relevant_context",
        "recover_interrupted_ingestions", "backfill_ready_file_metrics",
        "validate_document_file", "build_document_chunks", "update_file_ingestion",
        "insufficient_evidence", "buildRetrievePayload", "sourceLocationLabel",
        "openSource", "init_engine_stream", "_make_probe_pdf", "_run_probe",
        "preparePrivatePython", "prepareSeedModels", "startBackend",
        "waitForBackendHealth", "stopBackend", "XmuNativeBot", "OJClient",
        "asyncio.to_thread", "stream_sign", "is_active_time", "send_code",
        "_find_number_code", "send_radar", "build_prompt", "call_llm",
        "_strip_markdown_fence", "poll_result", "get_deadlines", "save_deadlines",
        "getStatusText", "getTimelineType", "getStatusClass", "ReadableStream",
        "TextDecoder", "EventSource", "KeepAlive", "SHA-256", "CSRF", "SSE",
        "MMR", "BM25", "CrossEncoder", "Chroma", "SQLite", "FastAPI",
        "PyMuPDF", "Tesseract", "Marker", "MinerU", "Electron", "Vue 3",
    )
    positions = []
    for term in code_terms:
        cursor = 0
        while True:
            idx = text.find(term, cursor)
            if idx < 0:
                break
            positions.append((idx, idx + len(term)))
            cursor = idx + len(term)
    positions.sort()
    merged = []
    for a, b in positions:
        if merged and a < merged[-1][1]:
            if b > merged[-1][1]:
                merged[-1] = (merged[-1][0], b)
        else:
            merged.append((a, b))
    for a, b in merged:
        if a > start:
            tokens.append((text[start:a], False))
        tokens.append((text[a:b], True))
        start = b
    if start < len(text):
        tokens.append((text[start:], False))
    if not tokens:
        tokens = [(text, False)]
    for value, emphasized in tokens:
        run = paragraph.add_run(value)
        format_run(run, bold=emphasized)


def add_body_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    add_inline_emphasis(p, text)
    p.paragraph_format.widow_control = True
    return p


def build():
    doc = Document()
    configure_document(doc)
    lines = CONTENT_PATH.read_text(encoding="utf-8").splitlines()

    first_content = True
    pending_page_break = False
    for raw in lines:
        if not raw.strip():
            continue
        if "]" not in raw or not raw.startswith("["):
            add_body_paragraph(doc, raw)
            continue
        marker, text = raw[1:].split("]", 1)
        text = text.strip()

        if marker == "COVER_TITLE":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(118)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            format_run(r, size=28, bold=True, color=DARK_BLUE)
            first_content = False
        elif marker == "COVER_SUBTITLE":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(30)
            r = p.add_run(text)
            format_run(r, size=18, bold=True, color=BLUE)
        elif marker == "META":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(text)
            format_run(r, size=10.5, color=MUTED)
        elif marker == "COVER_NOTE":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(30)
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(0)
            set_cell_shading(p, PALE_BLUE, BLUE)
            r = p.add_run(text)
            format_run(r, size=10.5, color=DARK_BLUE)
        elif marker == "PAGEBREAK":
            # Defer the break to the following major heading. A break-only
            # paragraph can create a blank page when the preceding section
            # already ends exactly at a page boundary.
            pending_page_break = True
        elif marker == "H1":
            p = doc.add_paragraph(text, style="Heading 1")
            if pending_page_break and (
                text == "使用说明"
                or text.startswith("第 1 页｜")
                or text.startswith("答辩追问速查")
                or text.startswith("复现清单")
            ):
                p.paragraph_format.page_break_before = True
            pending_page_break = False
            for run in p.runs:
                format_run(run, size=16, bold=True, color=BLUE)
        elif marker == "H2":
            p = doc.add_paragraph(text, style="Heading 2")
            for run in p.runs:
                format_run(run, size=13, bold=True, color=BLUE)
        elif marker == "H3":
            p = doc.add_paragraph(text, style="Heading 3")
            for run in p.runs:
                format_run(run, size=12, bold=True, color=DARK_BLUE)
        elif marker == "TIME":
            p = add_body_paragraph(doc, text, style="Time Label")
            set_cell_shading(p, PALE_BLUE, BLUE)
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.right_indent = Inches(0.1)
        elif marker == "CUE":
            p = add_body_paragraph(doc, text, style="Stage Cue")
            if any(word in text for word in ("过渡", "切到", "停顿")):
                p.paragraph_format.keep_with_next = False
                if len(doc.paragraphs) >= 2:
                    doc.paragraphs[-2].paragraph_format.keep_with_next = True
        elif marker == "P":
            add_body_paragraph(doc, text)
        elif marker == "NOTE":
            p = add_body_paragraph(doc, text)
            set_cell_shading(p, PALE_GRAY, DARK_BLUE)
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.right_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
        elif marker == "NUM":
            add_body_paragraph(doc, text, style="List Number")
        elif marker == "BULLET":
            add_body_paragraph(doc, text, style="List Bullet")
        else:
            add_body_paragraph(doc, text)

    # Remove the empty paragraph inserted by Document() if it remains before cover.
    if doc.paragraphs and not doc.paragraphs[0].text and len(doc.paragraphs[0]._p) == 1:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
